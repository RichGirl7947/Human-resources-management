from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


OUTPUT_PATH = Path(r"C:\Users\zhangliang\Desktop\孙敏笔记\孙敏-AI应用开发工程师-优化版.docx")

NAVY = "18324B"
TEAL = "0A8C82"
INK = "26384A"
MUTED = "66788A"
RULE = "B8C9D7"
WHITE = "FFFFFF"


def set_run_font(run, size=None, color=INK, bold=None, latin="Calibri") -> None:
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_bottom_border(paragraph, color=RULE, size=5, space=2) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_top_border(paragraph, color=RULE, size=4, space=4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), str(space))
    top.set(qn("w:color"), color)
    p_bdr.append(top)


def add_bullet_definition(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", "bullet"), ("w:lvlText", "\uf0b7"), ("w:lvlJc", "left")):
        node = OxmlElement(tag)
        node.set(qn("w:val"), value)
        lvl.append(node)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "340")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "340")
    ind.set(qn("w:hanging"), "170")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "28")
    spacing.set(qn("w:line"), "264")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Symbol")
    fonts.set(qn("w:hAnsi"), "Symbol")
    r_pr.append(fonts)
    lvl.append(r_pr)
    abstract.append(lvl)

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def add_hyperlink(paragraph, text: str, url: str, color=TEAL):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(r_fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "19")
    r_pr.append(size)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5.2)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=13.2, color=NAVY, bold=True)
    set_bottom_border(p)


def add_bullet(doc: Document, text: str, bullet_id: int, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.4)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(bullet_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=9.35, color=INK, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=9.35, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=9.35, color=INK)


def add_project_title(doc: Document, title: str, meta: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(1.6)
    p.paragraph_format.space_after = Pt(1.8)
    p.paragraph_format.line_spacing = 1.0
    r1 = p.add_run(title)
    set_run_font(r1, size=10.55, color=NAVY, bold=True)
    r2 = p.add_run(f" | {meta}")
    set_run_font(r2, size=9.75, color=TEAL, bold=True)


def add_body(doc: Document, text: str, size=9.55, color=INK, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.0)
    p.paragraph_format.line_spacing = 1.12
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=size, color=color, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=size, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(9.5)
    section.bottom_margin = Mm(10.5)
    section.left_margin = Mm(13)
    section.right_margin = Mm(13)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.12

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    set_top_border(fp)
    fr = fp.add_run("孙敏 | AI 应用开发工程师 | Agent 应用方向")
    set_run_font(fr, size=8.2, color=MUTED)


def build_resume() -> None:
    doc = Document()
    configure_document(doc)
    bullet_id = add_bullet_definition(doc)

    doc.core_properties.title = "孙敏 - AI 应用开发工程师简历"
    doc.core_properties.subject = "AI 应用开发 / Agent 工程"
    doc.core_properties.author = "孙敏"
    doc.core_properties.keywords = "Python, FastAPI, LangChain, Agent, RAG, PostgreSQL, Redis, Docker"

    header = doc.add_paragraph()
    header.paragraph_format.space_before = Pt(0)
    header.paragraph_format.space_after = Pt(2)
    header.paragraph_format.line_spacing = 1.0
    name = header.add_run("孙敏")
    set_run_font(name, size=23, color=NAVY, bold=True)
    role = header.add_run("   AI 应用开发工程师 | Agent 应用方向")
    set_run_font(role, size=13.6, color=TEAL, bold=True)

    contact = doc.add_paragraph()
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(1.2)
    contact.paragraph_format.line_spacing = 1.0
    cr = contact.add_run("电话：18808567947  |  邮箱：1931004473@qq.com  |  GitHub：")
    set_run_font(cr, size=9.55, color=INK)
    add_hyperlink(contact, "github.com/RichGirl7947", "https://github.com/RichGirl7947")

    education = doc.add_paragraph()
    education.paragraph_format.space_before = Pt(0)
    education.paragraph_format.space_after = Pt(3)
    education.paragraph_format.line_spacing = 1.0
    er = education.add_run("大连海洋大学 | 计算机科学与技术 | 本科 | 2025.07 毕业")
    set_run_font(er, size=9.45, color=MUTED)
    set_bottom_border(education, color=TEAL, size=7, space=3)

    add_section_heading(doc, "职业概述")
    add_body(
        doc,
        "具备 Python AI 应用开发与前端工程经验，能够基于 FastAPI、LangChain 和大模型 API 完成 Agent/RAG 应用、业务工具调用与 Web 工作台；重视结构化输出、人工复核、数据安全和可测试部署。",
        size=9.7,
    )

    add_section_heading(doc, "核心技能")
    add_bullet(doc, "Python 后端：Python、FastAPI、Pydantic、SQLAlchemy、PostgreSQL、REST API、pytest。", bullet_id, "Python 后端：")
    add_bullet(doc, "LLM / Agent：LangChain、Tool Calling、RAG、Qwen/OpenAI 兼容接口、Prompt 与离线评测。", bullet_id, "LLM / Agent：")
    add_bullet(doc, "工程部署：Redis/Celery、Docker Compose、Git/GitHub Actions；具备 Vue 3/TypeScript 前端能力。", bullet_id, "工程部署：")

    add_section_heading(doc, "项目经历")
    add_project_title(doc, "企业人力资源 Agent 平台", "独立项目 | 2026.08")
    add_bullet(
        doc,
        "基于 Python、FastAPI、LangChain 和 Vue 3 开发员工全生命周期平台，覆盖招聘、入职、员工服务、绩效和离职，七类业务能力以 Typed Tools 与 Supervisor 统一编排。",
        bullet_id,
    )
    add_bullet(
        doc,
        "实现职位画像、简历结构化评分和面试名单生成；HR 确认筛选后，在两天内通过 Celery 调度邮件/短信通知，并支持重试、幂等和失败补偿。",
        bullet_id,
    )
    add_bullet(
        doc,
        "实现 JWT/RBAC、Argon2、Fernet/HMAC 个人数据保护和审计日志；使用 PostgreSQL/Alembic 管理生产数据，Redis 作为任务队列与结果后端。",
        bullet_id,
    )
    add_bullet(
        doc,
        "提供 Docker Compose 与 Caddy HTTPS 部署配置；后端 9 项测试、Vue TypeScript 检查和生产构建均通过。",
        bullet_id,
    )

    add_project_title(doc, "Super Support 电商售后智能客服 Agent", "独立项目 | 2026.05 - 2026.06")
    project_link = doc.add_paragraph()
    project_link.paragraph_format.space_before = Pt(0)
    project_link.paragraph_format.space_after = Pt(1.5)
    project_link.paragraph_format.line_spacing = 1.0
    lr = project_link.add_run("项目地址：")
    set_run_font(lr, size=9.2, color=MUTED)
    add_hyperlink(project_link, "github.com/RichGirl7947/Intelligent-customer-service", "https://github.com/RichGirl7947/Intelligent-customer-service")
    add_bullet(
        doc,
        "基于 TypeScript、AI SDK 与 Qwen/OpenAI 兼容接口开发售后 Agent，将订单查询、退款校验、物流异常和人工升级封装为 8 项业务工具，以 Tool Calling 连接自然语言与确定性流程。",
        bullet_id,
    )
    add_bullet(
        doc,
        "由服务端注入 customerId，并通过 Guest RBAC、最小权限工具和退款二次确认降低事实编造、越权访问与高风险误操作；知识不足或工具失败时自动转人工。",
        bullet_id,
    )
    add_bullet(
        doc,
        "46 条 Vitest 与 10 条离线 Eval 全部通过；30 条真实 Qwen 场景中，26 次工具调用的参数 Schema 合法率与执行成功率均为 100%。",
        bullet_id,
    )

    add_section_heading(doc, "工作经历")
    add_project_title(doc, "杭州铭橙科技 | 前端工程师", "2025.02 - 2025.08")
    add_body(doc, "后台管理中心、App Web 端及生鲜微商城微信小程序", size=9.25, color=MUTED)
    add_bullet(
        doc,
        "参与内部管理系统与核心业务模块迭代，完成页面实现、接口联调、异常处理和日常维护；负责微信小程序的数据渲染、交互状态与异常反馈。",
        bullet_id,
    )
    add_bullet(
        doc,
        "封装基础 UI 与业务组件，使用资源压缩、路由懒加载和图片懒加载优化体验；使用 Node.js 完成接口 Mock，并参与 Code Review 与 Git Flow 协作。",
        bullet_id,
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_resume()
