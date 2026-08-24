from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT.parents[1] / "人力资源Agent项目介绍与技术总结.docx"
FLOWCHART_PATH = PROJECT_ROOT / "output" / "python_hr_agent_business_flow.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F7FB"
WHITE = "FFFFFF"
INK = "172033"
GOLD = "B7791F"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"表格列宽合计必须为 {TABLE_WIDTH_DXA} DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, color=INK, bold=None, italic=None, latin="Calibri") -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_header_footer(section, first_page=False) -> None:
    section.different_first_page_header_footer = first_page
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("PULSE HR  |  项目技术说明")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("PULSE HR  ·  ")
    set_run_font(fr, size=8.5, color=MUTED)
    add_page_number(fp)


def add_numbering_definition(doc: Document, num_format: str, level_text: str, font: str | None = None) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_format)
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), level_text)
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    if font:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), font)
        fonts.set(qn("w:hAnsi"), font)
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    # OOXML requires all abstractNum elements to precede num elements.
    # Inserting at the end makes Word ignore the custom definition and fall
    # back to a single continuing built-in list.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc: Document, text: str, num_id: int, bold_lead: str | None = None):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_body(doc: Document, text: str, bold_lead: str | None = None, italic=False):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)
    return p


def add_callout(doc: Document, label: str, text: str, fill=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, bold=True, color=DARK_BLUE)
    r2 = p.add_run(text)
    set_run_font(r2, color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA], indent_dxa=150)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "172033")
    set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=9, color="E7EEF8", latin="Consolas")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def style_table(table, header=True) -> None:
    table.style = "Table Grid"
    if header:
        set_repeat_table_header(table.rows[0])
        for cell in table.rows[0].cells:
            set_cell_shading(cell, LIGHT_BLUE)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=9.5, color=NAVY, bold=True)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    if row_index != 0 or not header:
                        set_run_font(run, size=9.2, color=INK)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(88)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PULSE HR  ·  TECHNICAL PROJECT BRIEF")
    set_run_font(r, size=10, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    tr = title.add_run("企业人力资源 Agent 平台")
    set_run_font(tr, size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    sr = subtitle.add_run("项目介绍、开发难点、技术栈与部署说明")
    set_run_font(sr, size=15, color=DARK_BLUE)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_after = Pt(60)
    lr = line.add_run("招聘  ·  入职  ·  员工服务  ·  绩效发展  ·  离职")
    set_run_font(lr, size=10.5, color=MUTED, bold=True)

    add_callout(
        doc,
        "项目定位",
        "基于 Python、FastAPI、Vue 3 与 LangChain 构建的员工全生命周期 HR Agent 系统；强调结构化输出、人工决策边界、个人数据保护与可靠消息调度。",
        fill=LIGHT_BLUE,
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(48)
    meta.paragraph_format.space_after = Pt(0)
    mr = meta.add_run(f"技术总结  |  {date.today().strftime('%Y-%m-%d')}")
    set_run_font(mr, size=10, color=MUTED)
    doc.add_page_break()


def add_project_snapshot(doc: Document, bullet_id: int) -> None:
    add_heading(doc, "执行摘要", 1)
    add_body(
        doc,
        "PULSE HR 是一个面向企业人力资源场景的 Agent 应用平台。项目以员工全生命周期为主线，将招聘、入职、在职服务、绩效发展和离职管理整合到同一套可视化工作台，并通过 LangChain 对七类专业能力进行统一编排。",
    )
    add_callout(
        doc,
        "一句话介绍",
        "这是一个让 Agent 负责分析与流程编排、让 HR 保留最终决策权的人力资源智能协作系统。",
    )

    add_heading(doc, "项目亮点", 2)
    highlights = [
        "完整业务闭环：从职位创建、简历筛选、面试通知，到员工入职、政策问答、绩效发展和离职归档。",
        "真实 Agent 框架：使用 LangChain LCEL、Typed Tools 和可选 Supervisor，而不是把大模型调用散落在业务代码中。",
        "Human in the Loop：录用、绩效、调岗、薪酬和离职等高影响事项必须由人工确认。",
        "生产化基础：提供 JWT/RBAC、Argon2、字段加密、审计日志、PostgreSQL、Alembic、Redis、Celery 和 HTTPS 配置。",
        "可演示、可测试：无模型 Key 时仍可运行确定性 Agent；配置百炼后可启用 qwen-flash Supervisor。",
    ]
    for item in highlights:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "核心业务模块", 2)
    table = doc.add_table(rows=1, cols=3)
    headers = ["业务阶段", "Agent 能力", "人工控制点"]
    for i, value in enumerate(headers):
        table.cell(0, i).text = value
    rows = [
        ("招聘", "职位画像、简历匹配、面试名单与通知调度", "职位审批、完成筛选、录用决定"),
        ("入职", "生成 HR、IT、经理与员工任务清单", "资料核验与入职确认"),
        ("在职服务", "政策问答、请求识别与审批链生成", "低置信度转人工、流程审批"),
        ("绩效发展", "按绩效结果生成发展或改进建议", "经理与 HR 最终评审"),
        ("离职", "生成跨部门交接与结算清单", "离职审批与最终归档"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_geometry(table, [1600, 4760, 3000])
    style_table(table)


def add_business_and_architecture(doc: Document, bullet_id: int) -> None:
    add_heading(doc, "一、项目介绍与业务架构", 1)
    add_heading(doc, "1.1 建设背景", 2)
    add_body(
        doc,
        "传统 HR 系统通常聚焦表单录入和记录查询，而招聘筛选、制度咨询、流程路由、发展建议等工作仍依赖人工反复判断。本项目希望把这些重复、可结构化的分析工作交给 Agent，同时将高风险决策明确留给 HR 与业务负责人。",
    )
    add_heading(doc, "1.2 设计目标", 2)
    for item in [
        "业务统一：用一套平台覆盖员工全生命周期，而不是构建彼此割裂的功能页面。",
        "能力可解释：Agent 返回分数、匹配依据、缺口、来源和执行轨迹。",
        "决策可控：Agent 只给出建议、清单和路由，不替代最终人事决定。",
        "系统可运行：未配置大模型时保留确定性规则，便于开发、测试和故障降级。",
        "数据可保护：个人信息加密存储，写操作保留不含请求正文的审计记录。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "1.3 员工全生命周期流程", 2)
    if FLOWCHART_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        picture = p.add_run().add_picture(str(FLOWCHART_PATH), width=Inches(6.5))
        picture._inline.docPr.set("title", "人力资源 Agent 业务流程图")
        picture._inline.docPr.set(
            "descr",
            "展示招聘、入职、在职服务、绩效发展与离职管理五个业务阶段，以及 Agent 自动化节点、人工审批节点和流程走向。",
        )
        caption = doc.add_paragraph("图 1  人力资源 Agent 业务流程图")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(2)
        caption.paragraph_format.space_after = Pt(8)
        for run in caption.runs:
            set_run_font(run, size=9, color=MUTED, italic=True)

    add_heading(doc, "1.4 技术调用路径", 2)
    add_callout(
        doc,
        "主链路",
        "Vue 3 可视化页面 → FastAPI REST API → LangChain LCEL / Tools → 业务 Agent → SQLAlchemy → SQLite（开发）或 PostgreSQL（生产）",
    )
    add_callout(
        doc,
        "异步链路",
        "完成简历筛选 → 创建面试通知 → Celery / Redis 定时调度 → SMTP 邮件与短信 Webhook → 状态回写与失败补偿",
        fill=LIGHT_GRAY,
    )


def add_agent_design(doc: Document) -> None:
    add_heading(doc, "二、Agent 架构设计", 1)
    add_heading(doc, "2.1 七类 Agent 能力", 2)
    table = doc.add_table(rows=1, cols=3)
    for i, value in enumerate(["Agent", "主要输入", "结构化输出"]):
        table.cell(0, i).text = value
    rows = [
        ("职位画像", "职位、部门、职责、技能、薪资、学历、经验", "岗位使命、核心职责、技能、筛选问题、成功指标"),
        ("简历筛选", "目标职位、技能要求、简历文本", "匹配分、命中技能、技能缺口、推荐状态"),
        ("入职引导", "员工、部门、职位", "HR/IT/经理/员工任务清单"),
        ("政策问答", "政策问题", "答案、依据、置信度、人工升级标识"),
        ("流程路由", "请假、报销、调岗等自然语言请求", "流程类别、审批人链路、下一步"),
        ("绩效教练", "绩效分数、经理反馈", "发展类别与行动建议"),
        ("离职交接", "离职原因", "交接、资产、账号、证明与结算清单"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_geometry(table, [1750, 3540, 4070])
    style_table(table)

    add_heading(doc, "2.2 LangChain 的使用方式", 2)
    add_body(doc, "LCEL 统一路由：", bold_lead="LCEL 统一路由：")
    add_body(doc, "业务接口不直接依赖某个模型，而是调用统一的 LangChainRuntime。运行名称、标签、Agent 名称和人工复核策略都作为调用元数据传入，便于后续接入链路追踪。")
    add_body(doc, "Typed Tools：", bold_lead="Typed Tools：")
    add_body(doc, "七类能力注册为带类型定义的工具，Supervisor 可以根据用户意图自动选择工具，工具内部仍返回统一的 AgentResult。")
    add_body(doc, "可选 Supervisor：", bold_lead="可选 Supervisor：")
    add_body(doc, "配置阿里云百炼后，ChatOpenAI 通过 OpenAI 兼容协议连接 qwen-flash；未配置 Key 时，确定性 Agent 仍可独立运行。")
    add_body(doc, "Human in the Loop：", bold_lead="Human in the Loop：")
    add_body(doc, "系统提示词明确禁止模型替代录用、绩效定级、薪酬、调岗和离职审批，结构化结果中也携带 human_review_required 标志。")


def add_challenges(doc: Document, number_id: int) -> None:
    add_heading(doc, "三、开发过程中最难解决的问题", 1)
    add_body(
        doc,
        "本项目真正困难的部分并不是页面数量，而是如何让 Agent、业务规则、安全边界和异步流程共同工作。以下问题最能体现项目的工程含量。",
    )
    challenges = [
        (
            "让 Agent 可控，而不是只会聊天",
            "直接让大模型输出招聘或绩效结论容易产生格式漂移、幻觉和越权决策。",
            "将能力拆成七个单一职责 Agent，统一返回 AgentResult；使用 Typed Tools 和 LCEL 编排；对高影响结果强制标记人工复核。",
            "模型负责理解和建议，业务状态变化仍由明确 API 与人工按钮控制。",
        ),
        (
            "真实百炼模型的兼容接入",
            "API Key、Provider、模型名和 Base URL 需要保持一致，且不能把真实 Key 写进代码或示例文件。",
            "通过 Settings 从系统环境变量读取配置，使用 LangChain ChatOpenAI 的 OpenAI 兼容模式接入百炼，并提供安全配置脚本。",
            "模型提供方与业务逻辑解耦，未配置模型时自动回退到确定性运行层。",
        ),
        (
            "AI 筛选前五名与两天内通知",
            "排序、重复执行、通知计划、邮件与短信双通道、失败恢复必须保持一致；简单地在接口中直接发送消息会导致超时和重复发送。",
            "完成筛选时按匹配度生成最多五人的面试名单，提前创建两种通知记录，默认 24 小时后调度；Celery 使用延迟确认、失败重试、指数退避和定时扫描补偿。",
            "接口快速返回，通知状态可追踪，服务重启后仍可继续处理未完成任务。",
        ),
        (
            "个人数据既要加密，又要支持去重查询",
            "如果只加密邮箱，每次密文不同，无法通过数据库唯一索引判断重复；如果只哈希，又无法展示原始数据。",
            "使用 Fernet 保存可解密字段，使用基于密钥的 HMAC 摘要完成邮箱或登录标识的去重查询；审计日志不记录请求正文。",
            "同时满足展示、查询和数据最小暴露要求。",
        ),
        (
            "开发环境与生产数据库的一致性",
            "本地 SQLite 便于演示，但生产需要 PostgreSQL；两者在锁、约束、连接池和迁移方式上存在差异。",
            "SQLAlchemy 统一 ORM，开发环境允许自动建表，生产环境强制通过 Alembic upgrade head；就绪检查执行真实数据库查询。",
            "本地开发门槛低，生产结构又能被版本化管理。",
        ),
        (
            "登录初始化与角色权限的闭环",
            "首次管理员、开发环境初始化令牌、密码规则和页面缓存状态容易让用户误以为系统无法登录。",
            "增加初始化状态接口，仅在生产显示初始化令牌；前端将 Pydantic 校验错误转换为可读提示；JWT 依赖统一校验角色权限。",
            "管理员初始化、登录、账号创建和权限拒绝都有明确状态。",
        ),
        (
            "把安全能力放进业务代码而不破坏可用性",
            "JWT、Argon2、字段加密、可信主机、CSP、HSTS、审计日志和跨域策略涉及多个层次，遗漏任何一层都可能造成上线风险。",
            "安全配置集中化，FastAPI 中间件统一添加安全响应头和写操作审计；生产环境启动时校验关键配置，Caddy 负责 HTTPS。",
            "安全要求成为默认行为，而不是部署前临时补丁。",
        ),
    ]
    for index, (title, difficulty, solution, result) in enumerate(challenges, start=1):
        add_list_item(doc, f"{title}", number_id, bold_lead=title)
        add_body(doc, f"难点：{difficulty}", bold_lead="难点：")
        add_body(doc, f"解决方案：{solution}", bold_lead="解决方案：")
        add_body(doc, f"结果与取舍：{result}", bold_lead="结果与取舍：")


def add_tech_stack(doc: Document, bullet_id: int) -> None:
    add_heading(doc, "四、技术栈", 1)
    table = doc.add_table(rows=1, cols=3)
    for i, value in enumerate(["层次", "技术", "项目用途"]):
        table.cell(0, i).text = value
    rows = [
        ("前端", "Vue 3、TypeScript、Vite", "可视化工作台、表单、状态管理、角色化页面"),
        ("后端", "Python 3.11+、FastAPI、Pydantic 2、Uvicorn", "REST API、依赖注入、数据校验、应用服务"),
        ("Agent", "LangChain 1.x、LCEL、Typed Tools、create_agent", "Agent 编排、工具调用、Supervisor 和结构化结果"),
        ("大模型", "阿里云百炼 qwen-flash", "通过 OpenAI 兼容协议支持自然语言工具选择"),
        ("数据访问", "SQLAlchemy 2", "ORM、事务、索引、外键和 PostgreSQL 行锁"),
        ("数据库", "SQLite、PostgreSQL 17、Alembic", "本地演示、生产存储、结构迁移"),
        ("认证安全", "JWT、Argon2、RBAC", "登录、密码哈希、四类角色权限"),
        ("数据保护", "Fernet、HMAC", "个人数据字段加密与可查询摘要"),
        ("异步任务", "Celery 5、Redis 7", "定时通知、失败重试、任务补偿"),
        ("外部通知", "SMTP、短信 Webhook", "面试邮件与短信发送"),
        ("部署", "Docker、Docker Compose、Caddy", "多阶段构建、服务编排、HTTPS 反向代理"),
        ("测试", "pytest、TestClient、monkeypatch、vue-tsc", "接口、Agent、权限、加密、消息幂等与类型检查"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_geometry(table, [1500, 3100, 4760])
    style_table(table)

    add_heading(doc, "4.1 体现的 Python 知识", 2)
    for item in [
        "面向对象与抽象：BaseAgent 抽象基类、多个专业 Agent 实现和注册表。",
        "类型系统：Annotated、泛型容器、联合类型、Pydantic 模型和返回类型。",
        "装饰器与依赖注入：FastAPI 路由、Celery 任务、LangChain tool、字段校验和缓存。",
        "异步编程：应用生命周期、async/await、后台任务和线程切换。",
        "上下文管理与事务：数据库 Session、SMTP 连接以及提交、回滚边界。",
        "安全编程：密码哈希、JWT、加解密、HMAC、常量时间令牌比较。",
        "测试技术：fixture、TestClient、monkeypatch、幂等性测试和权限测试。",
    ]:
        add_list_item(doc, item, bullet_id)


def add_deployment(doc: Document, number_id: int, bullet_id: int) -> None:
    add_heading(doc, "五、部署方式", 1)
    add_heading(doc, "5.1 本地开发部署", 2)
    add_body(doc, "本地模式使用 SQLite，适合功能开发、演示和自动化测试。")
    add_code_block(
        doc,
        [
            "python -m venv .venv",
            ".\\.venv\\Scripts\\Activate.ps1",
            "python -m pip install -e \".[dev]\"",
            "uvicorn hr_agent.main:app --reload",
        ],
    )
    add_body(doc, "访问地址：http://127.0.0.1:8000。首次打开时创建管理员，之后通过工号和密码登录。")

    add_heading(doc, "5.2 生产部署拓扑", 2)
    table = doc.add_table(rows=1, cols=3)
    for i, value in enumerate(["服务", "职责", "关键机制"]):
        table.cell(0, i).text = value
    rows = [
        ("Caddy", "公网入口与反向代理", "自动 HTTPS、压缩、安全响应头"),
        ("FastAPI Web", "页面与业务 API", "JWT/RBAC、审计、就绪检查"),
        ("PostgreSQL", "业务数据存储", "持久卷、健康检查、Alembic"),
        ("Redis", "消息代理与结果后端", "密码、AOF 持久化"),
        ("Celery Worker", "执行异步通知", "延迟确认、失败重试"),
        ("Celery Beat", "扫描和补偿任务", "周期调度未完成通知"),
        ("db-tools", "备份与恢复", "pg_dump / pg_restore"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_geometry(table, [1900, 3560, 3900])
    style_table(table)

    add_heading(doc, "5.3 生产部署步骤", 2)
    steps = [
        "准备 Linux 服务器和域名，将 DNS 指向服务器，并开放 80、443 端口。",
        "复制 .env.production.example 为 .env.production，填写域名、数据库密码、Redis 密码、JWT 密钥、初始化令牌和 Fernet 密钥。",
        "根据需要配置百炼、SMTP 和短信 Webhook；真实密钥不得提交到代码仓库。",
        "执行 docker compose --env-file .env.production up -d --build。",
        "Web 容器启动时先执行 alembic upgrade head，再启动 Uvicorn。",
        "检查 /health 与 /health/ready，确认数据库、页面、登录、队列和 HTTPS 均正常。",
        "完成首位管理员初始化后，安全保存或移除一次性初始化令牌。",
        "执行备份并进行一次恢复演练，确认灾难恢复链路有效。",
    ]
    for step in steps:
        add_list_item(doc, step, number_id)
    add_code_block(
        doc,
        [
            "Copy-Item .env.production.example .env.production",
            "# 编辑 .env.production，填写真实域名和密钥",
            "docker compose --env-file .env.production up -d --build",
            "docker compose ps",
        ],
    )

    add_heading(doc, "5.4 关键环境变量", 2)
    vars_table = doc.add_table(rows=1, cols=2)
    vars_table.cell(0, 0).text = "配置项"
    vars_table.cell(0, 1).text = "用途"
    var_rows = [
        ("HR_ENVIRONMENT", "生产环境设为 production"),
        ("HR_DATABASE_URL", "PostgreSQL 连接串"),
        ("HR_JWT_SECRET", "JWT 签名密钥，至少 32 字符"),
        ("HR_DATA_ENCRYPTION_KEY", "Fernet 个人数据加密密钥"),
        ("HR_BOOTSTRAP_TOKEN", "首位管理员初始化令牌"),
        ("HR_CELERY_BROKER_URL", "Redis 消息代理地址"),
        ("DASHSCOPE_API_KEY", "阿里云百炼 API Key"),
        ("HR_SMTP_* / HR_SMS_WEBHOOK_URL", "面试邮件与短信通道"),
        ("HR_DOMAIN / HR_ALLOWED_HOSTS", "HTTPS 域名与可信主机"),
    ]
    for row in var_rows:
        cells = vars_table.add_row().cells
        cells[0].text, cells[1].text = row
    set_table_geometry(vars_table, [3400, 5960])
    style_table(vars_table)


def add_security_quality(doc: Document, bullet_id: int) -> None:
    add_heading(doc, "六、安全、可靠性与质量保障", 1)
    add_heading(doc, "6.1 登录与权限", 2)
    for item in [
        "Argon2 保存密码哈希，明文密码不进入数据库。",
        "JWT 访问令牌存储于浏览器 sessionStorage，默认 60 分钟过期。",
        "RBAC 支持管理员、HR、招聘人员和只读人员四类角色。",
        "连续五次登录失败后锁定账号 15 分钟。",
        "首次生产初始化要求一次性令牌。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "6.2 个人数据保护", 2)
    for item in [
        "候选人和员工姓名、邮箱、手机号、简历、筛选画像、HR 请求、绩效反馈和离职原因使用 Fernet 加密。",
        "需要去重查询的标识使用 HMAC 摘要，不保存可逆的明文索引。",
        "写操作审计只保存操作者、路径、状态码、来源 IP 和 User-Agent，不复制请求正文。",
        "生产环境缺少 JWT、数据加密密钥、PostgreSQL 或 Celery Broker 时拒绝启动。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "6.3 消息可靠性", 2)
    for item in [
        "数据库先保存通知记录，再提交 Celery 任务，业务状态和任务状态可追踪。",
        "Worker 使用 task_acks_late 与 task_reject_on_worker_lost，降低进程异常造成的消息丢失。",
        "发送失败最多重试八次，使用指数退避并限制最长等待时间。",
        "Beat 周期扫描到期、失败和长时间 Sending 状态，形成补偿机制。",
        "短信携带 idempotency_key；邮件使用稳定 Message-ID，降低重复通知风险。",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "6.4 测试与当前状态", 2)
    status_table = doc.add_table(rows=1, cols=2)
    status_table.cell(0, 0).text = "检查项"
    status_table.cell(0, 1).text = "最近结果"
    status_rows = [
        ("后端 pytest", "9 项测试通过"),
        ("Python compileall", "通过"),
        ("Vue TypeScript 与生产构建", "通过"),
        ("前端生产依赖审计", "未发现漏洞"),
        ("全新 Alembic 迁移", "升级到 Head，未发现模型差异"),
        ("Docker Compose 配置", "语法有效"),
        ("运行接口", "健康检查 200、未登录业务接口 401、管理员接口 200"),
        ("加密落库", "当前管理员标识与姓名均为密文"),
    ]
    for row in status_rows:
        cells = status_table.add_row().cells
        cells[0].text, cells[1].text = row
    set_table_geometry(status_table, [3600, 5760])
    style_table(status_table)


def add_limits_and_interview(doc: Document, bullet_id: int) -> None:
    add_heading(doc, "七、当前边界与后续改进", 1)
    add_callout(
        doc,
        "当前定位",
        "项目已经具备完整 MVP 与生产化基础，但当前运行实例仍是本地开发环境；PostgreSQL、Redis、Celery、真实邮件/短信和公网 HTTPS 需要在生产服务器上完成联调。",
        fill="FFF8E8",
    )
    improvements = [
        "升级存在安全公告的 Python 依赖，并在 CI 中持续执行依赖审计。",
        "增加修改密码、管理员重置密码、登录限流、MFA 或企业 SSO。",
        "把系统账号内部字段从历史 email 命名正式迁移为 employee_number。",
        "增加 PostgreSQL、Redis、Celery、SMTP、短信和 Caddy 的集成测试。",
        "引入 LangSmith 或 OpenTelemetry，补齐 Agent 链路追踪、成本和时延监控。",
        "建设企业制度 RAG 知识库、评测集、提示词版本管理和模型降级策略。",
        "为审计日志、备份、密钥轮换和个人数据保留建立运营制度。",
    ]
    for item in improvements:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "八、面试介绍参考", 1)
    add_heading(doc, "90 秒项目介绍", 2)
    add_callout(
        doc,
        "参考表达",
        "我开发了一个企业人力资源 Agent 平台，覆盖招聘、入职、员工服务、绩效和离职五个阶段。后端使用 Python、FastAPI 和 SQLAlchemy，前端使用 Vue 3；Agent 层通过 LangChain LCEL、Typed Tools 和 Supervisor 统一编排，并接入阿里云百炼 qwen-flash。项目的重点不是让模型直接做人事决定，而是让 Agent 完成职位画像、简历匹配、政策回答和流程路由，同时通过 Human in the Loop 保留 HR 的最终审批权。为了接近真实生产环境，我还实现了 JWT/RBAC、Argon2、Fernet 字段加密、审计日志、PostgreSQL/Alembic、Redis/Celery、失败重试、Docker Compose 和 Caddy HTTPS。开发中最困难的是保证 Agent 输出可控、通知不丢失，以及在加密个人数据的同时支持唯一查询。",
    )
    add_heading(doc, "建议重点展开的问题", 2)
    for item in [
        "为什么不用一个大 Agent，而是拆成七个专业 Agent？",
        "如何保证 Agent 不替代 HR 做最终录用或绩效决定？",
        "Celery、Redis、数据库通知记录如何共同保证可靠发送？",
        "当前项目距离真正生产上线还差哪些能力？",
    ]:
        add_list_item(doc, item, bullet_id)

def main() -> None:
    doc = Document()
    for section in doc.sections:
        configure_page(section)
        configure_header_footer(section, first_page=True)
    configure_styles(doc)
    doc.core_properties.title = "企业人力资源 Agent 平台项目介绍与技术总结"
    doc.core_properties.subject = "项目介绍、开发难点、技术栈与部署说明"
    doc.core_properties.author = "PULSE HR 项目组"
    doc.core_properties.keywords = "Python, FastAPI, Vue3, LangChain, HR Agent, PostgreSQL, Celery"

    bullet_id = add_numbering_definition(doc, "bullet", "\uf0b7", font="Symbol")
    challenge_number_id = add_numbering_definition(doc, "decimal", "%1.")
    deployment_number_id = add_numbering_definition(doc, "decimal", "%1.")

    add_cover(doc)
    add_project_snapshot(doc, bullet_id)
    add_business_and_architecture(doc, bullet_id)
    add_agent_design(doc)
    add_challenges(doc, challenge_number_id)
    add_tech_stack(doc, bullet_id)
    add_deployment(doc, deployment_number_id, bullet_id)
    add_security_quality(doc, bullet_id)
    add_limits_and_interview(doc, bullet_id)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
